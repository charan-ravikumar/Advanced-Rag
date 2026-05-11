import os
from pathlib import Path

import fitz
import pandas as pd
import frontmatter

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from markdown_it import MarkdownIt

from schemas import Document, Section


# ============================================
# PDF LOADER
# ============================================

def load_pdf(file_path):

    doc = fitz.open(file_path)

    full_text = ""
    sections = []

    for page_num, page in enumerate(doc):

        text = page.get_text()

        full_text += text + "\n"

        sections.append(
            Section(
                content=text,
                section_title=f"Page {page_num + 1}",
                metadata={
                    "page": page_num + 1
                }
            )
        )

    return Document(
        content=full_text,
        metadata={
            "source": file_path,
            "file_type": ".pdf"
        },
        sections=sections
    )


# ============================================
# DOCX LOADER
# ============================================

def load_docx(file_path):

    doc = DocxDocument(file_path)

    sections = []

    current_heading = "Introduction"
    current_content = []

    for para in doc.paragraphs:

        style = para.style.name.lower()

        if "heading" in style:

            if current_content:

                sections.append(
                    Section(
                        content="\n".join(current_content).strip(),
                        section_title=current_heading
                    )
                )

            current_heading = para.text
            current_content = []

        else:

            if para.text.strip():
                current_content.append(para.text)

    if current_content:

        sections.append(
            Section(
                content="\n".join(current_content).strip(),
                section_title=current_heading
            )
        )

    full_text = "\n".join(
        [section.content for section in sections]
    )

    return Document(
        content=full_text,
        metadata={
            "source": file_path,
            "file_type": ".docx"
        },
        sections=sections
    )


# ============================================
# HTML LOADER
# ============================================

def load_html(file_path):

    with open(file_path, "r", encoding="utf-8") as f:
        html = f.read()

    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()

    sections = []

    headers = soup.find_all(["h1", "h2", "h3"])

    for header in headers:

        content = []

        for sibling in header.find_next_siblings():

            if sibling.name in ["h1", "h2", "h3"]:
                break

            text = sibling.get_text(" ", strip=True)

            if text:
                content.append(text)

        sections.append(
            Section(
                content="\n".join(content).strip(),
                section_title=header.get_text(strip=True)
            )
        )

    full_text = soup.get_text(separator="\n")

    return Document(
        content=full_text,
        metadata={
            "source": file_path,
            "file_type": ".html"
        },
        sections=sections
    )


# ============================================
# MARKDOWN LOADER
# ============================================

def load_md(file_path):

    # ----------------------------------------
    # LOAD FRONTMATTER + CONTENT
    # ----------------------------------------

    post = frontmatter.load(file_path)

    md_text = post.content

    frontmatter_metadata = post.metadata

    # ----------------------------------------
    # PARSE MARKDOWN
    # ----------------------------------------

    md = MarkdownIt()

    tokens = md.parse(md_text)

    sections = []

    current_heading = "Introduction"

    current_content = []

    expecting_heading = False

    # ----------------------------------------
    # TOKEN WALK
    # ----------------------------------------

    for token in tokens:

        # Detect heading start
        if token.type == "heading_open":

            # Save previous section
            if current_content:

                sections.append(
                    Section(
                        content="\n".join(current_content).strip(),
                        section_title=current_heading
                    )
                )

            current_content = []

            expecting_heading = True

        # Capture heading text
        elif token.type == "inline" and expecting_heading:

            current_heading = token.content

            expecting_heading = False

        # Capture normal content
        elif token.type == "inline":

            if token.content.strip():
                current_content.append(token.content)

    # ----------------------------------------
    # FINAL SECTION
    # ----------------------------------------

    if current_content:

        sections.append(
            Section(
                content="\n".join(current_content).strip(),
                section_title=current_heading
            )
        )

    # ----------------------------------------
    # RETURN DOCUMENT
    # ----------------------------------------

    return Document(
        content=md_text,
        metadata={
            "source": file_path,
            "file_type": ".md",
            **frontmatter_metadata
        },
        sections=sections
    )


# ============================================
# TXT LOADER
# ============================================

def load_txt(file_path):

    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

    return Document(
        content=text,
        metadata={
            "source": file_path,
            "file_type": ".txt"
        },
        sections=[
            Section(
                content=text,
                section_title="Raw Text"
            )
        ]
    )


# ============================================
# XLSX LOADER
# ============================================

def load_xlsx(file_path):

    excel = pd.ExcelFile(file_path)

    sections = []

    full_text = ""

    for sheet_name in excel.sheet_names:

        df = pd.read_excel(
            file_path,
            sheet_name=sheet_name
        )

        text = df.to_string(index=False)

        full_text += text + "\n"

        sections.append(
            Section(
                content=text,
                section_title=sheet_name
            )
        )

    return Document(
        content=full_text,
        metadata={
            "source": file_path,
            "file_type": ".xlsx"
        },
        sections=sections
    )


# ============================================
# LOADER MAP
# ============================================

LOADER_MAP = {
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".html": load_html,
    ".md": load_md,
    ".txt": load_txt,
    ".xlsx": load_xlsx,
}


# ============================================
# ROUTER
# ============================================

def load_document(file_path):

    extension = Path(file_path).suffix.lower()

    loader = LOADER_MAP.get(extension)

    if not loader:
        raise ValueError(
            f"Unsupported file type: {extension}"
        )

    return loader(file_path)